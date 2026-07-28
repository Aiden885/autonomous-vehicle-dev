#!/usr/bin/env python3
"""Generate weekly project report docx from a small Markdown input file."""

from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass, field
from pathlib import Path


DEFAULT_OWNER = "赵煜坤"
DEFAULT_FONT = "微软雅黑"
SEPARATOR = "——————————————————————————————"


@dataclass
class ProgressGroup:
    title: str
    rows: list[tuple[str, str, str]] = field(default_factory=list)


@dataclass
class ReportData:
    meta: dict[str, str] = field(default_factory=dict)
    progress: list[ProgressGroup] = field(default_factory=list)
    risks: list[tuple[str, str, str]] = field(default_factory=list)
    next_plans: list[tuple[str, str, str]] = field(default_factory=list)
    coordination: list[tuple[str, str, str]] = field(default_factory=list)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a GAASD weekly report docx.")
    parser.add_argument("input", type=Path, help="Markdown weekly report input.")
    parser.add_argument("-o", "--output", type=Path, help="Output docx path.")
    return parser.parse_args()


def parse_front_matter(lines: list[str]) -> tuple[dict[str, str], list[str]]:
    meta: dict[str, str] = {}
    if not lines or lines[0].strip() != "---":
        return meta, lines

    end_index = None
    for index in range(1, len(lines)):
        if lines[index].strip() == "---":
            end_index = index
            break

    if end_index is None:
        return meta, lines

    for raw in lines[1:end_index]:
        if ":" not in raw:
            continue
        key, value = raw.split(":", 1)
        meta[key.strip()] = value.strip().strip('"')

    return meta, lines[end_index + 1 :]


def section_name(title: str) -> tuple[str, str]:
    clean = title.strip().lstrip("#").strip()
    for sep in ("：", ":"):
        if sep in clean:
            key, value = clean.split(sep, 1)
            return key.strip().lower(), value.strip()
    return clean.lower(), clean


def split_row(line: str) -> list[str]:
    content = line.strip()
    if content.startswith("- "):
        content = content[2:].strip()
    return [part.strip() for part in content.split("|")]


def parse_report(path: Path) -> ReportData:
    text = path.read_text(encoding="utf-8")
    meta, body_lines = parse_front_matter(text.splitlines())
    report = ReportData(meta=meta)
    owner = meta.get("owner", DEFAULT_OWNER)

    mode = ""
    current_group: ProgressGroup | None = None

    for raw in body_lines:
        line = raw.strip()
        if not line or line.startswith("<!--"):
            continue

        if line.startswith("## "):
            key, value = section_name(line)
            current_group = None
            if key in {"进展", "progress", "核心任务进展"}:
                mode = "progress"
                group_title = value if value != key else "核心任务进展"
                current_group = ProgressGroup(group_title)
                report.progress.append(current_group)
            elif "风险" in key:
                mode = "risks"
            elif "下周" in key or "计划" in key:
                mode = "next"
            elif "协调" in key or "帮助" in key:
                mode = "coordination"
            else:
                mode = ""
            continue

        if line.startswith("- "):
            cols = split_row(line)
            if mode == "progress" and report.progress:
                group = report.progress[-1]
                if len(cols) >= 3:
                    group.rows.append((cols[0], cols[1], " | ".join(cols[2:])))
                elif len(cols) == 2:
                    group.rows.append((cols[0], owner, cols[1]))
                elif len(cols) == 1:
                    group.rows.append((cols[0], owner, cols[0]))
            elif mode == "risks":
                if len(cols) >= 3:
                    report.risks.append((cols[0], cols[1], " | ".join(cols[2:])))
                elif len(cols) == 2:
                    report.risks.append((cols[0], "", cols[1]))
                elif len(cols) == 1:
                    report.risks.append((cols[0], "", "持续跟进"))
            elif mode == "next":
                if len(cols) >= 3:
                    report.next_plans.append((cols[0], cols[1], " | ".join(cols[2:])))
                elif len(cols) == 2:
                    report.next_plans.append((cols[0], owner, cols[1]))
                elif len(cols) == 1:
                    report.next_plans.append((cols[0], owner, cols[0]))
            elif mode == "coordination":
                if len(cols) >= 3:
                    report.coordination.append((cols[0], cols[1], " | ".join(cols[2:])))
                elif len(cols) == 2:
                    report.coordination.append((cols[0], owner, cols[1]))
                elif len(cols) == 1:
                    report.coordination.append((cols[0], owner, cols[0]))

    return report


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None) -> None:
    run.font.name = DEFAULT_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc, text: str, size: float = 11, bold: bool = False, center: bool = False):
    paragraph = doc.add_paragraph()
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=size, bold=bold)
    return paragraph


def add_table(doc, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        paragraph = header_cells[index].paragraphs[0]
        run = paragraph.add_run(header)
        set_run_font(run, size_pt=10.5, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            run = paragraph.add_run(str(value))
            set_run_font(run, size_pt=10.5, bold=False)

    doc.add_paragraph()


def configure_doc(doc) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = DEFAULT_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    normal.font.size = Pt(10.5)


def output_path(input_path: Path, report: ReportData, explicit: Path | None) -> Path:
    if explicit is not None:
        return explicit
    title = report.meta.get("title", input_path.stem)
    safe_title = title.replace("/", "_").replace("\\", "_")
    return Path("docs") / f"{safe_title}.docx"


def generate_docx(report: ReportData, output: Path) -> None:
    doc = Document()
    configure_doc(doc)

    title = report.meta.get("title", "项目进展周报")
    report_period = report.meta.get("report_period", "")
    next_period = report.meta.get("next_period", "")

    add_paragraph(doc, title, size=18, bold=True, center=True)
    add_paragraph(doc, "项目进展周报", size=14, bold=True, center=True)
    if report_period:
        add_paragraph(doc, f"报告周期：{report_period}", size=11, center=True)
    add_paragraph(doc, SEPARATOR, size=10.5)

    add_paragraph(doc, "一、核心任务进展", size=13, bold=True)
    if not report.progress:
        add_table(doc, ["任务项", "负责人", "完成情况"], [("本周工作整理", DEFAULT_OWNER, "待补充")])
    else:
        for index, group in enumerate(report.progress, start=1):
            add_paragraph(doc, f"{index}. {group.title}", size=11.5, bold=True)
            rows = group.rows or [("本项工作", report.meta.get("owner", DEFAULT_OWNER), "待补充")]
            add_table(doc, ["任务项", "负责人", "完成情况"], rows)

    add_paragraph(doc, SEPARATOR, size=10.5)
    add_paragraph(doc, "二、风险与问题", size=13, bold=True)
    risks = report.risks or [("暂无阻塞性问题", "当前无新增阻塞", "持续跟踪 GAASD 版本和联调结果")]
    add_table(doc, ["问题/风险", "影响", "应对措施"], risks)

    add_paragraph(doc, SEPARATOR, size=10.5)
    add_paragraph(doc, "三、下周工作计划", size=13, bold=True)
    if next_period:
        add_paragraph(doc, f"计划周期：{next_period}", size=10.5)
    next_plans = report.next_plans or [("下周工作", report.meta.get("owner", DEFAULT_OWNER), "待补充")]
    add_table(doc, ["计划项", "负责人", "计划内容"], next_plans)

    if report.coordination:
        add_paragraph(doc, SEPARATOR, size=10.5)
        add_paragraph(doc, "四、需要协调与帮助", size=13, bold=True)
        add_table(doc, ["事项", "负责人", "说明"], report.coordination)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    args = parse_args()
    if not args.input.exists():
        print(f"输入文件不存在: {args.input}", file=sys.stderr)
        return 1

    report = parse_report(args.input)
    out = output_path(args.input, report, args.output)
    generate_docx(report, out)
    print(out)
    return 0


try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError as exc:
    if exc.name == "docx":
        print("缺少依赖 python-docx，请先执行: python3 -m pip install python-docx", file=sys.stderr)
        raise SystemExit(2)
    raise


if __name__ == "__main__":
    raise SystemExit(main())
